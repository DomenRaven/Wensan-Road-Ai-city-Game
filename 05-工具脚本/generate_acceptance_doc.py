# -*- coding: utf-8 -*-
"""Generate AI学习小游戏创作-功能验收文档.docx from project truth sources."""
from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT = Path(__file__).resolve().parents[1]
OUT = ROOT / "AI学习小游戏创作-功能验收文档.docx"

ITEMS: list[tuple[str, str | None, list[tuple[str, str, str, str]]]] = [
    (
        "一、系统导航与公共功能",
        None,
        [
            (
                "1.1",
                "A 链 / B 链入口",
                "展厅提供快玩与教育创作两条入口",
                "A 链 `/kiosk/` 可进入 S0–S9 快玩；B 链 `/kiosk/edu/` 可进入 B0–B7 教育创作；页头可返回 A 链",
            ),
            (
                "1.2",
                "蓝白主题界面",
                "B 链与 A 链采用统一蓝白 K12 展厅 UI",
                "页面背景浅蓝白、卡片白底、主按钮品牌蓝；代码区浅灰底；**不**按品类整页换肤（C-03 决策）",
            ),
            (
                "1.3",
                "横竖屏自适应",
                "Kiosk 根据视口自动切换横屏/竖屏布局",
                "宽度 breakpoint 900px；竖屏单栏步骤与双栏 grid 切换正常；`body[data-orientation]` 正确",
            ),
            (
                "1.4",
                "触控尺寸约束",
                "按钮与卡片满足展厅触控最小尺寸",
                "主要按钮/选项卡高度 ≥48px；配方选项卡 ≥72px（`kiosk_edu_spec.touch_constraints`）",
            ),
            (
                "1.5",
                "一键启动工坊",
                "提供本地一键启动 API 与 Kiosk 静态服务",
                "运行 `启动游戏工坊.exe` 或 `launch_workshop.py` 后，`:8000` API 与 `:8080` Kiosk 可访问；B 链须从 `http://127.0.0.1:8080/kiosk/edu/` 进入",
            ),
            (
                "1.6",
                "系统触控软键盘",
                "文本输入唤起 Windows TabTip，点外/导航自动收起",
                "B1/B2 `edu-touch-input` 点击调用 `POST /kiosk/touch-keyboard/show`；点输入区外、上一步/下一步 `pointerdown` 调用 `hide`；`warm` 预加载；`tabtip_native` COM 线程",
            ),
            (
                "1.7",
                "页眉今日榜单",
                "B6/B7 阶段页眉提供日榜入口",
                "「今日榜单」钮与「← 展厅快玩」「重新开始」同高；点击打开 `leaderboard.js` 霓虹弹层；七款 Tab 可切换",
            ),
        ],
    ),
    (
        "二、A 链 · 展厅快玩",
        None,
        [
            (
                "2.1",
                "A 链入口与导航",
                "A 链页头展示「展厅快玩」模式与作品名区域",
                "打开 `/kiosk/` 可见 GameForge K12 品牌、进度点、上一步/下一步/重新开始",
            ),
            (
                "2.2",
                "S0 作品起名",
                "快玩流程可选为游戏起名",
                "S0 步骤可输入作品名并写入会话 `display_name`",
            ),
            (
                "2.3",
                "S1 品类选择",
                "七款游戏大卡片选类",
                "S1 展示 platformer、shmup、survivor、pingpong、fighting、parkour、racing 共 7 款可选",
            ),
            (
                "2.4",
                "S9 直启试玩",
                "选定品类后直接启动 Godot 试玩",
                "完成 S1 后调用 `POST .../play/launch` 启动外置 Godot 窗口，无代码区与证书流程",
            ),
            (
                "2.5",
                "跳转 B 链",
                "快玩页提供进入教育创作入口",
                "页脚链接「创作工坊（B 链教育版 B0–B7）」可进入 `/kiosk/edu/`",
            ),
        ],
    ),
    (
        "三、B 链 · 准备与创作阶段",
        "准备与连接",
        [
            (
                "3.1.1",
                "B0 开馆连接",
                "进入 B 链后自动创建会话并校验环境",
                "展示欢迎加载动画；`GET /bootstrap` 通过后进入 B1；失败时阻断并提示",
            ),
            (
                "3.1.2",
                "会话持久化",
                "会话 ID 本地存储并可恢复",
                "刷新页面后会话可续（Redis 或内存存储）；`GET /sessions/{id}` 可读取状态",
            ),
        ],
    ),
    (
        "三、B 链 · 准备与创作阶段",
        "B1 意图匹配",
        [
            (
                "3.2.1",
                "自然语言输入",
                "儿童描述想玩的游戏类型",
                "B1 文本框可输入；点击下一步调用 `POST /intent/match-genre` 匹配 7 品类之一",
            ),
            (
                "3.2.2",
                "关键词兜底",
                "API 不可用时本地关键词匹配",
                "断网或 API 失败时使用 `intent_genre_lexicon` 关键词 fallback，仍能进入 B2",
            ),
            (
                "3.2.3",
                "肥皂泡泡快捷选类",
                "B1 全屏彩色泡泡代替原椭圆示例按钮",
                "7 个品类泡泡可拖拽碰撞；点击泡泡将文案填入输入框；`prefers-reduced-motion` 时静态网格降级",
            ),
        ],
    ),
    (
        "三、B 链 · 准备与创作阶段",
        "B2 起名与揭幕",
        [
            (
                "3.3.1",
                "填写你的名字",
                "B2 第一屏询问「你的名字是？」",
                "须先输入 1–8 字 creator_name；校验通过后 `PATCH /sessions/{id}` 持久化；触控键盘可正常输入中文",
            ),
            (
                "3.3.2",
                "游戏起名",
                "B2 第二屏为游戏起名与推荐名",
                "可手动输入或点推荐名芯片；写入 `display_name`；顶栏显示「名字《游戏名》」",
            ),
            (
                "3.3.3",
                "B2 子步骤导航",
                "上一步在名字屏与游戏名屏间切换",
                "游戏名屏点上一步回到名字屏；名字屏点上一步回到 B1",
            ),
            (
                "3.3.4",
                "揭幕卡与烟花",
                "B2 完成后播放揭幕过渡再进入 B3",
                "`EduForgeReveal.play` 展示「开始制作{creator}的{game}游戏！」文案与烟花动画；结束后进入双栏 B3",
            ),
        ],
    ),
    (
        "三、B 链 · 准备与创作阶段",
        "B3–B4 工作区与配方",
        [
            (
                "3.4.1",
                "B3 双栏工作区",
                "左侧代码区、右侧品类预览/引导",
                "进入 B3 后切换双栏布局；左栏文件树+代码阅读器；右栏品类 HTML 预览动画",
            ),
            (
                "3.4.2",
                "B4 逐卡轮播",
                "配方问卷一题一卡展示",
                "`EduB4CardFlow` 每道 tuning 题单独卡片；选答后切下一卡；末卡「完成配方」可提交",
            ),
            (
                "3.4.3",
                "B4 进度条",
                "顶部显示答题进度",
                "进度条随 `confirmedIndices` 更新；全部作答后进度满格",
            ),
            (
                "3.4.4",
                "B4 题面 widget",
                "按品类展示互动小部件",
                "各品类卡片含 thumb/slider 等 widget（见 `b4-card-flow.js`）；与题面选项联动",
            ),
            (
                "3.4.5",
                "B4 卡内上一步",
                "配方阶段可在卡片间回退",
                "B4 步骤点「上一步」调用 `EduB4CardFlow.prev` 回到上一题；首卡上一步回到 B3",
            ),
            (
                "3.4.6",
                "B4 配方提交",
                "用户选择手感选项并提交",
                "每题须作答；提交 `POST .../creative/answers` 成功；**无**小技能勾选题（`q_skill` 已移除）",
            ),
            (
                "3.4.7",
                "配方数值合并",
                "答案写入 workspace 的 game_config",
                "generate 后 tuning 在 ±30% clamp 内；试玩体感与所选选项可感知差异",
            ),
        ],
    ),
    (
        "四、B 链 · 制作阶段",
        "需求分析与代码生成",
        [
            (
                "4.1.1",
                "B5 需求分析",
                "提交配方后进行需求分析（preset 路径）",
                "`POST .../analyze-requirements` 返回 resolutions，resolution 为 preset；**未接入 LLM 实时生成**",
            ),
            (
                "4.1.2",
                "B5 多文件代码剧场",
                "模拟 AI 逐文件写出代码",
                "按 `edu_workspace_trees.json` 多文件自上而下打字展示；背景浅底；可跳过等待",
            ),
            (
                "4.1.3",
                "B5 品类剧场 overlay",
                "代码打字时叠加品类趣味动效",
                "`EduTheaterOverlays` 按七款 genre 注入半透明 overlay（见 `theater-overlays/`）；打字结束 `stop` 清理",
            ),
            (
                "4.1.4",
                "generate/v2 工作区生成",
                "复制模板到隔离工作区并合并配置",
                "`POST .../generate/v2` 生成 `workspace/{session_id}/`；返回 `code_map` 供高亮；不修改 `templates/` 原目录",
            ),
            (
                "4.1.5",
                "B6 真代码展示",
                "制作完成后展示 workspace 内真实文件",
                "可读 `game_config.json`；文件树可点开 `.gd`/`.json`/`.tscn` 节选；代码区纵向滚动可读",
            ),
        ],
    ),
    (
        "五、B 链 · 试玩与教育联动",
        "作品证书",
        [
            (
                "5.1.1",
                "证书趣味标题",
                "证书主标题使用创作者与游戏名",
                "标题格式 `{creator}的{game}！`（如「小明的星星大冒险！」）；非通用「作品登记证书」作主标题",
            ),
            (
                "5.1.2",
                "B6 证书自动展示",
                "制作完成时展示作品登记证书",
                "含品类中文名、创作时间、配方摘要表（≥3 行有效问答）",
            ),
            (
                "5.1.3",
                "证书配方摘要",
                "摘要来自 B4 答案与模板题面",
                "每道已答 tuning 题显示「问题 + 你的选择」；不显示已删除的小技能题",
            ),
            (
                "5.1.4",
                "证书 PNG 保存",
                "一键导出证书图片到本机",
                "点击「保存证书」`#btnCertSave`；`html2canvas` 渲染 `#edu-cert-card` 并触发浏览器下载 PNG",
            ),
            (
                "5.1.5",
                "证书扫码下载",
                "证书含二维码供手机下载同图",
                "后端签发 token；`GET /public/certificates/{token}` 返回 PNG；**展馆公网域名未配时扫码仅 API 就绪**",
            ),
            (
                "5.1.6",
                "B7 复看证书",
                "试玩阶段可再次打开证书",
                "B7 工具栏「查看证书」可重新打开证书叠层",
            ),
        ],
    ),
    (
        "五、B 链 · 试玩与教育联动",
        "Godot 试玩与代码高亮",
        [
            (
                "5.2.1",
                "B6/B7 启动试玩",
                "一键启动 Godot 外置试玩窗口",
                "`POST .../play/launch` 成功返回 pid；状态卡显示已启动",
            ),
            (
                "5.2.2",
                "Godot 窗口归位",
                "Windows 下自动贴边至引导区",
                "横屏贴屏幕右半与 `godot-zone` 对齐；竖屏贴主显示器下半；失败时提示手动找窗且 launch 不阻断",
            ),
            (
                "5.2.3",
                "游戏动作上报",
                "试玩操作写入教育动作日志",
                "各品类 `_edu/{slug}_hooks.gd` 写 `.edu_actions.jsonl`；前端 `GET .../play/actions` 轮询",
            ),
            (
                "5.2.4",
                "B7 代码行高亮",
                "操作触发左侧代码黄框高亮与解说气泡",
                "根据 `code_map` 与 `code_anchors` 定位行号；caption 气泡显示操作说明",
            ),
            (
                "5.2.5",
                "讲解员演示钮",
                "未真实操作时也可触发教学高亮",
                "B6/B7 右栏按品类展示演示钮（如 platformer：跳/踩怪/捡金币）；点击模拟 `play/action`",
            ),
            (
                "5.2.6",
                "试玩进程状态",
                "轮询 Godot 是否仍在运行",
                "`GET .../play/status` 返回 running true/false；关窗后 running 变为 false",
            ),
        ],
    ),
    (
        "五、B 链 · 试玩与教育联动",
        "日榜排行榜",
        [
            (
                "5.3.1",
                "关窗触发日榜",
                "试玩结束关 Godot 窗后弹出排行榜",
                "platformer / shmup / survivor 等品类关窗后 `leaderboard.js` 展示本局得分与当日榜单",
            ),
            (
                "5.3.2",
                "七款 Tab 切换",
                "榜单弹层可按品类查看",
                "弹层顶部 7 个 Tab（横版闯关、街机飞机…）；切换后 `GET .../leaderboard/daily/{genre}` 刷新列表",
            ),
            (
                "5.3.3",
                "乒乓球单局规则",
                "乒乓球仅在整局结束结算上榜",
                "`pingpong_hooks` 发 `session_end` 才弹榜；单分得分不触发弹层",
            ),
            (
                "5.3.4",
                "日榜持久化",
                "当日成绩本机 JSON 存储",
                "按 `Asia/Shanghai` 自然日；Redis 可用时同步；断网可用 `edu_leaderboard_session_fallback` 降级",
            ),
            (
                "5.3.5",
                "页眉随时查榜",
                "不试玩也可从页眉打开榜单",
                "B6/B7 页眉「今日榜单」打开同一弹层；「继续」关闭后回到当前步骤",
            ),
        ],
    ),
    (
        "六、后端会话与工作区",
        "会话管理",
        [
            (
                "6.1.1",
                "健康检查",
                "API 服务可用性探测",
                "`GET /health` 返回正常",
            ),
            (
                "6.1.2",
                "会话池",
                "多路并发会话隔离",
                "默认最多 10 路会话；每会话独立 `workspace/{session_id}/`",
            ),
            (
                "6.1.3",
                "Redis 降级",
                "Redis 不可用时内存存储",
                "无 Redis 时会话仍可创建与续用（重启 API 后丢失）",
            ),
            (
                "6.1.4",
                "会话释放",
                "离开页面或重新开始时释放资源",
                "`POST .../release` 或 pagehide sendBeacon；bootstrap 可清理孤立 workspace",
            ),
        ],
    ),
    (
        "七、七款游戏模板",
        "模板与品类",
        [
            (
                "7.1.1",
                "横版闯关 platformer",
                "横版跳跃闯关模板可生成可试玩",
                "模板路径 `templates/platformer/`；B4 共 5 道 tuning 题",
            ),
            (
                "7.1.2",
                "街机飞机射击 shmup",
                "纵向射击模板可生成可试玩",
                "模板路径 `templates/shmup/`；B4 共 5 道 tuning 题",
            ),
            (
                "7.1.3",
                "生存升级 survivor",
                "割草生存模板可生成可试玩",
                "模板路径 `templates/survivor/`；B4 共 5 道 tuning 题",
            ),
            (
                "7.1.4",
                "乒乓球 pingpong",
                "双人对打乒乓球模板可生成可试玩",
                "模板路径 `templates/pingpong/`；B4 共 5 道 tuning 题",
            ),
            (
                "7.1.5",
                "格斗对战 fighting",
                "横版格斗模板可生成可试玩",
                "模板路径 `templates/fighting/`；B4 共 4 道 tuning 题",
            ),
            (
                "7.1.6",
                "跑酷 parkour",
                "无尽跑酷模板可生成可试玩",
                "模板路径 `templates/parkour/`；B4 共 5 道 tuning 题",
            ),
            (
                "7.1.7",
                "欢乐赛车 racing",
                "赛车模板可生成可试玩",
                "模板路径 `templates/racing/`；B4 共 5 道 tuning 题",
            ),
        ],
    ),
    (
        "八、展厅触屏试玩（触控 overlay）",
        "无键盘触屏操控",
        [
            (
                "8.1.1",
                "platformer 虚拟键",
                "横版闯关触屏操控",
                "注入 `platformer_touch_overlay.gd`：左下 ← → hold，右下跳跃；按钮 ≥48px",
            ),
            (
                "8.1.2",
                "parkour 虚拟键",
                "跑酷触屏跳跃与滑铲",
                "注入 `parkour_touch_overlay.gd`：跳跃与滑铲虚拟键可完成一局",
            ),
            (
                "8.1.3",
                "survivor 摇杆瞄准",
                "割草生存触屏移动与射击",
                "注入 `survivor_touch_overlay.gd`：虚拟摇杆移动 + 触区瞄准射击",
            ),
            (
                "8.1.4",
                "fighting 虚拟键",
                "格斗触屏方向与拳脚",
                "注入 `fighting_touch_overlay.gd`：方向、轻拳、重拳、格挡、大招虚拟键",
            ),
            (
                "8.1.5",
                "pingpong 触屏拖动",
                "乒乓球拍上下控制",
                "注入 `pingpong_touch_overlay.gd`：触屏拖动映射 move_up/move_down",
            ),
            (
                "8.1.6",
                "shmup 触屏拖动",
                "战机横向跟随触屏位置",
                "触屏拖动控制战机 X 轴（沿用 `player_ship.gd` 鼠标跟随语义）；可射击游玩",
            ),
            (
                "8.1.7",
                "racing 触屏转向",
                "赛车左右滑动转向",
                "注入 `racing_touch_overlay.gd`：底部水平滑动映射 steer_left/right",
            ),
        ],
    ),
]


def set_cell_shading(cell, fill: str = "D9E8F5") -> None:
    from docx.oxml import OxmlElement

    tc = cell._tc
    tc_pr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def add_table(doc: Document, rows: list[tuple[str, str, str, str]]) -> None:
    table = doc.add_table(rows=1, cols=5)
    table.style = "Table Grid"
    headers = ("编号", "功能项", "功能描述", "验收标准", "验收结果")
    hdr = table.rows[0].cells
    for i, text in enumerate(headers):
        hdr[i].text = text
        for p in hdr[i].paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.size = Pt(10)
        set_cell_shading(hdr[i], "B4C7E7")

    for rid, name, desc, criteria in rows:
        row = table.add_row().cells
        row[0].text = rid
        row[1].text = name
        row[2].text = desc
        row[3].text = criteria
        row[4].text = "□通过    □不通过"
        for c in row:
            for p in c.paragraphs:
                for r in p.runs:
                    r.font.size = Pt(10)

    doc.add_paragraph()


def build() -> None:
    doc = Document()

    for section in doc.sections:
        section.top_margin = Cm(2)
        section.bottom_margin = Cm(2)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("AI学习小游戏创作 · 功能验收文档")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0x1E, 0x3A, 0x5F)

    doc.add_paragraph()

    meta_lines = [
        "项目名称：文三路 AI 游戏创作工坊（GameForge K12）",
        "产品版本：1.2（7.4 P4-B/C 展厅落地）· Git tag v1.2 待用户确认",
        "文档版本：v1.2 · 2026-07-06 · 对齐 P4-E2E 收工基线",
        "验收范围：本文档仅包含当前仓库已实现的功能点验收，不含规划未开发项。",
        "不含范围：大模型实时代码生成、Godot 浏览器内嵌（P4-A 挂起）、联机/存档/内购/商城、"
        "展馆公网扫码域名（C-08 API 已就绪，手机实机待配）。",
        "验收说明：逐项操作验证，在「验收结果」栏勾选通过或不通过，并在备注栏填写说明。",
        "访问地址：Kiosk http://127.0.0.1:8080/kiosk/edu/ · API http://127.0.0.1:8000/docs",
        "生成脚本：05-工具脚本/generate_acceptance_doc.py",
    ]
    for line in meta_lines:
        p = doc.add_paragraph(line)
        for r in p.runs:
            r.font.size = Pt(11)

    doc.add_paragraph()

    current_section = ""
    for section_title, group_title, rows in ITEMS:
        if section_title != current_section:
            current_section = section_title
            h = doc.add_heading(section_title, level=1)
            for r in h.runs:
                r.font.color.rgb = RGBColor(0x1E, 0x40, 0xAF)

        if group_title:
            gh = doc.add_heading(group_title, level=2)
            for r in gh.runs:
                r.font.size = Pt(13)

        add_table(doc, rows)

    doc.add_heading("验收签字", level=1)
    sign_table = doc.add_table(rows=4, cols=4)
    sign_table.style = "Table Grid"
    sign_headers = ["角色", "姓名", "日期", "签字"]
    for i, h in enumerate(sign_headers):
        sign_table.rows[0].cells[i].text = h
        set_cell_shading(sign_table.rows[0].cells[i], "B4C7E7")
    for label in ("开发方", "测试方", "甲方验收"):
        row = sign_table.add_row().cells
        row[0].text = label

    doc.save(OUT)
    print(f"Wrote {OUT}")


if __name__ == "__main__":
    build()
