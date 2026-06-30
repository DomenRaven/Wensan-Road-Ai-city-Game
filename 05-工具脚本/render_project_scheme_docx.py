"""将 开发文档/项目方案_v1.1.md 渲染为排版美观的 Word 文档。"""
from __future__ import annotations

import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

ROOT: Path = Path(__file__).resolve().parent.parent
SRC: Path = ROOT / "开发文档" / "项目方案_v1.1.md"
OUT: Path = ROOT / "开发文档" / "项目方案_v1.1.docx"

ACCENT: RGBColor = RGBColor(0x25, 0x63, 0xEB)
TEXT: RGBColor = RGBColor(0x0F, 0x17, 0x2A)
MUTED: RGBColor = RGBColor(0x64, 0x74, 0x8B)
HEADER_BG: str = "DBEAFE"
CODE_BG: str = "F1F5F9"
FONT_BODY: str = "微软雅黑"
FONT_CODE: str = "Consolas"


def set_cell_shading(cell: object, fill_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()  # type: ignore[attr-defined]
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill_hex)
    shd.set(qn("w:val"), "clear")
    tc_pr.append(shd)


def set_run_font(run: object, name: str, size_pt: float, bold: bool = False, color: RGBColor | None = None) -> None:
    run.font.name = name  # type: ignore[attr-defined]
    run._element.rPr.rFonts.set(qn("w:eastAsia"), name)  # type: ignore[attr-defined]
    run.font.size = Pt(size_pt)  # type: ignore[attr-defined]
    run.font.bold = bold  # type: ignore[attr-defined]
    if color is not None:
        run.font.color.rgb = color  # type: ignore[attr-defined]


def add_body_paragraph(doc: Document, text: str, *, italic: bool = False, muted: bool = False) -> None:
    para = doc.add_paragraph()
    para.paragraph_format.space_after = Pt(6)
    para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    para.paragraph_format.line_spacing = 1.25
    run = para.add_run(text)
    set_run_font(run, FONT_BODY, 11, color=MUTED if muted else TEXT)
    run.italic = italic  # type: ignore[attr-defined]


def add_code_block(doc: Document, lines: list[str]) -> None:
    table = doc.add_table(rows=1, cols=1)
    table.autofit = True
    cell = table.rows[0].cells[0]
    set_cell_shading(cell, CODE_BG)
    cell.text = ""
    for idx, line in enumerate(lines):
        para = cell.paragraphs[0] if idx == 0 else cell.add_paragraph()
        para.paragraph_format.space_after = Pt(0)
        para.paragraph_format.line_spacing = 1.15
        run = para.add_run(line)
        set_run_font(run, FONT_CODE, 9, color=TEXT)
    doc.add_paragraph()


def parse_table_row(line: str) -> list[str]:
    parts = [p.strip() for p in line.strip().strip("|").split("|")]
    return parts


def is_table_sep(line: str) -> bool:
    return bool(re.match(r"^\|\s*:?-+:?\s*(\|\s*:?-+:?\s*)+\|?\s*$", line.strip()))


def strip_md_inline(text: str) -> str:
    text = re.sub(r"\*\*(.+?)\*\*", r"\1", text)
    text = re.sub(r"`([^`]+)`", r"\1", text)
    text = re.sub(r"\[(.+?)\]\([^)]+\)", r"\1", text)
    return text


def add_table(doc: Document, rows: list[list[str]]) -> None:
    if not rows:
        return
    col_count = max(len(r) for r in rows)
    table = doc.add_table(rows=len(rows), cols=col_count)
    table.style = "Table Grid"
    for r_idx, row in enumerate(rows):
        for c_idx in range(col_count):
            cell = table.rows[r_idx].cells[c_idx]
            value = strip_md_inline(row[c_idx]) if c_idx < len(row) else ""
            cell.text = ""
            para = cell.paragraphs[0]
            para.paragraph_format.space_after = Pt(2)
            run = para.add_run(value)
            is_header = r_idx == 0
            set_run_font(run, FONT_BODY, 10 if not is_header else 10.5, bold=is_header, color=TEXT if is_header else TEXT)
            if is_header:
                set_cell_shading(cell, HEADER_BG)
    doc.add_paragraph()


def render_markdown(md_path: Path, out_path: Path) -> None:
    lines: list[str] = md_path.read_text(encoding="utf-8").splitlines()
    doc = Document()

    section = doc.sections[0]
    section.top_margin = Cm(2.2)
    section.bottom_margin = Cm(2.0)
    section.left_margin = Cm(2.5)
    section.right_margin = Cm(2.5)

    i = 0
    title_done = False
    in_code = False
    code_buf: list[str] = []
    table_buf: list[list[str]] = []

    def flush_table() -> None:
        nonlocal table_buf
        if table_buf:
            add_table(doc, table_buf)
            table_buf = []

    while i < len(lines):
        raw = lines[i]
        line = raw.rstrip()

        if line.strip().startswith("```"):
            if in_code:
                add_code_block(doc, code_buf)
                code_buf = []
                in_code = False
            else:
                flush_table()
                in_code = True
            i += 1
            continue

        if in_code:
            code_buf.append(line)
            i += 1
            continue

        if not line.strip():
            flush_table()
            i += 1
            continue

        if line.strip() == "---":
            flush_table()
            doc.add_paragraph()
            i += 1
            continue

        if line.startswith("|") and is_table_sep(line):
            i += 1
            continue

        if line.startswith("|"):
            table_buf.append(parse_table_row(line))
            i += 1
            continue

        flush_table()

        if line.startswith("# ") and not title_done:
            title = strip_md_inline(line[2:].strip())
            para = doc.add_paragraph()
            para.alignment = WD_ALIGN_PARAGRAPH.CENTER
            para.paragraph_format.space_after = Pt(12)
            run = para.add_run(title)
            set_run_font(run, FONT_BODY, 22, bold=True, color=ACCENT)
            title_done = True
            i += 1
            continue

        if line.startswith("## "):
            text = strip_md_inline(line[3:].strip())
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(14)
            para.paragraph_format.space_after = Pt(8)
            run = para.add_run(text)
            set_run_font(run, FONT_BODY, 16, bold=True, color=ACCENT)
            i += 1
            continue

        if line.startswith("### "):
            text = strip_md_inline(line[4:].strip())
            para = doc.add_paragraph()
            para.paragraph_format.space_before = Pt(10)
            para.paragraph_format.space_after = Pt(6)
            run = para.add_run(text)
            set_run_font(run, FONT_BODY, 13, bold=True, color=TEXT)
            i += 1
            continue

        if line.startswith("> "):
            add_body_paragraph(doc, strip_md_inline(line[2:].strip()), muted=True)
            i += 1
            continue

        if line.startswith("- "):
            para = doc.add_paragraph(style="List Bullet")
            para.paragraph_format.space_after = Pt(3)
            run = para.add_run(strip_md_inline(line[2:].strip()))
            set_run_font(run, FONT_BODY, 11, color=TEXT)
            i += 1
            continue

        if line.strip().startswith("*") and line.strip().endswith("*"):
            add_body_paragraph(doc, strip_md_inline(line.strip().strip("*")), italic=True, muted=True)
            i += 1
            continue

        add_body_paragraph(doc, strip_md_inline(line.strip()))
        i += 1

    flush_table()
    if in_code and code_buf:
        add_code_block(doc, code_buf)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(out_path))


def main() -> int:
    src = SRC
    out = OUT
    if len(sys.argv) >= 2:
        src = Path(sys.argv[1])
    if len(sys.argv) >= 3:
        out = Path(sys.argv[2])
    if not src.is_file():
        print(f"源文件不存在: {src}", file=sys.stderr)
        return 1
    render_markdown(src, out)
    print(f"OK: {out}")
    return 0


if __name__ == "__main__":
    raise SystemExit(main())
